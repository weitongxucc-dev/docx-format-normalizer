#!/usr/bin/env python3
"""
DOCX Format Normalizer - Output Validator
Checks key formatting properties of a processed .docx file against a template.

Usage:
    python3 validate.py --input output.docx --template templates/government_document.json
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is not installed. Run: bash scripts/setup.sh")
    sys.exit(1)


class FormatValidator:
    def __init__(self, docx_path, template_path):
        self.doc = Document(docx_path)
        with open(template_path, "r", encoding="utf-8") as f:
            self.template = json.load(f)
        self.passed = 0
        self.failed = 0
        self.warnings = 0
        self.results = []

    def run(self):
        self._check_page_setup()
        self._check_body_text()
        self._check_styles()
        self._check_headings()
        self._check_tables()
        self._check_force_clear()
        self._check_content_compliance()
        self._check_cover_page()
        self._report()

    def _check(self, name, condition, detail_ok, detail_fail):
        if condition:
            self.passed += 1
            self.results.append(("PASS", name, detail_ok))
        else:
            self.failed += 1
            self.results.append(("FAIL", name, detail_fail))

    def _warn(self, name, detail):
        self.warnings += 1
        self.results.append(("WARN", name, detail))

    def _check_page_setup(self):
        ps = self.template.get("page_setup")
        if not ps:
            return

        section = self.doc.sections[0]
        margins = ps.get("margins", {})

        if ps.get("paper_size") == "A4":
            w_mm = round(section.page_width / 36000, 1)
            h_mm = round(section.page_height / 36000, 1)
            self._check("Paper Size",
                        abs(w_mm - 210) < 1 and abs(h_mm - 297) < 1,
                        f"A4 confirmed ({w_mm}x{h_mm}mm)",
                        f"Expected A4 (210x297mm), got {w_mm}x{h_mm}mm")

        for key, attr in [("top_mm", "top_margin"), ("bottom_mm", "bottom_margin"),
                          ("left_mm", "left_margin"), ("right_mm", "right_margin")]:
            if key in margins:
                expected = margins[key]
                actual_mm = round(getattr(section, attr) / 36000, 1)
                self._check(f"Margin {key}",
                            abs(actual_mm - expected) < 1,
                            f"{actual_mm}mm (expected {expected}mm)",
                            f"Got {actual_mm}mm, expected {expected}mm")

    def _check_body_text(self):
        body = self.template.get("body_text")
        if not body:
            return

        expected_font = body.get("font_cn")
        expected_size = body.get("font_size_pt")
        expected_align = body.get("alignment")

        first_heading_idx = None
        for i, para in enumerate(self.doc.paragraphs):
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name or "title" in style_name:
                first_heading_idx = i
                break

        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name or "title" in style_name:
                continue

            if first_heading_idx is not None and i < first_heading_idx:
                continue

            if para.runs:
                run = para.runs[0]
                r = run._element
                rPr = r.find(qn('w:rPr'))
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None and expected_font:
                        actual_ea = rFonts.get(qn('w:eastAsia'))
                        self._check("Body Font (eastAsia)",
                                    actual_ea == expected_font,
                                    f"{actual_ea} (expected {expected_font})",
                                    f"Got {actual_ea}, expected {expected_font}")
                break

        try:
            style = self.doc.styles['Normal']
            if expected_size:
                actual_size = style.font.size
                if actual_size:
                    actual_pt = actual_size.pt
                    self._check("Normal Style Font Size",
                                abs(actual_pt - expected_size) < 1,
                                f"{actual_pt}pt (expected {expected_size}pt)",
                                f"Got {actual_pt}pt, expected {expected_size}pt")
        except (KeyError, Exception):
            self._warn("Normal Style", "Could not read Normal style")

    def _check_styles(self):
        body = self.template.get("body_text", {})
        line_rule = body.get("line_spacing_rule")

        try:
            style = self.doc.styles['Normal']
            pf = style.paragraph_format

            if line_rule == "fixed":
                from docx.enum.text import WD_LINE_SPACING
                self._check("Normal Line Spacing Rule",
                            pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY,
                            "EXACTLY (correct for fixed)",
                            f"Expected EXACTLY, got {pf.line_spacing_rule}")

                sb = pf.space_before
                sa = pf.space_after
                sb_pt = sb.pt if sb else 0
                sa_pt = sa.pt if sa else 0
                self._check("Normal Space Before (fixed mode)",
                            sb_pt == 0,
                            f"{sb_pt}pt (expected 0)",
                            f"Got {sb_pt}pt, expected 0")
                self._check("Normal Space After (fixed mode)",
                            sa_pt == 0,
                            f"{sa_pt}pt (expected 0)",
                            f"Got {sa_pt}pt, expected 0")
        except (KeyError, Exception):
            self._warn("Style Check", "Could not read Normal style")

    def _check_headings(self):
        headings = self.template.get("headings", [])
        if not headings:
            return

        for h in headings:
            level = h.get("level", 0)
            if level == 0:
                style_name = "Title"
            else:
                style_name = f"Heading {level}"

            try:
                style = self.doc.styles[style_name]
                expected_font = h.get("font_cn", h.get("font", ""))
                expected_bold = h.get("bold")

                if expected_font:
                    rpr = style.element.find(qn('w:rPr'))
                    if rpr is not None:
                        rFonts = rpr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            actual = rFonts.get(qn('w:eastAsia'))
                            self._check(f"{style_name} Font",
                                        actual == expected_font,
                                        f"{actual} (expected {expected_font})",
                                        f"Got {actual}, expected {expected_font}")

                if expected_bold is not None:
                    actual_bold = style.font.bold
                    self._check(f"{style_name} Bold",
                                actual_bold == expected_bold,
                                f"bold={actual_bold} (expected {expected_bold})",
                                f"Got bold={actual_bold}, expected {expected_bold}")
            except KeyError:
                self._warn(f"{style_name}", "Style not found in document")

    def _check_tables(self):
        table_cfg = self.template.get("tables")
        if not table_cfg or not self.doc.tables:
            return

        body = self.template.get("body_text", {})
        expected_cell_font = table_cfg.get("cell_font_cn", body.get("font_cn", ""))

        table = self.doc.tables[0]
        if table.rows and table.rows[0].cells:
            cell = table.rows[0].cells[0]
            for para in cell.paragraphs:
                if para.runs:
                    run = para.runs[0]
                    r = run._element
                    rPr = r.find(qn('w:rPr'))
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None and expected_cell_font:
                            actual = rFonts.get(qn('w:eastAsia'))
                            self._check("Table Cell Font",
                                        actual == expected_cell_font,
                                        f"{actual} (expected {expected_cell_font})",
                                        f"Got {actual}, expected {expected_cell_font}")
                    break

    def _check_force_clear(self):
        clear_rules = self.template.get("force_clear", [])
        if not clear_rules:
            return

        black = RGBColor(0, 0, 0)
        bold_found = False
        colored_found = False

        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.bold:
                    bold_found = True
                if run.font.color and run.font.color.rgb and run.font.color.rgb != black:
                    colored_found = True

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.bold:
                                bold_found = True

        if "bold" in clear_rules:
            self._check("Force Clear: Bold",
                        not bold_found,
                        "No bold found (correct)",
                        "Bold still present (should be cleared)")

        if "colored_text" in clear_rules:
            self._check("Force Clear: Colors",
                        not colored_found,
                        "No colored text (correct)",
                        "Colored text still present")

        for section in self.doc.sections:
            header = section.header
            footer = section.footer
            has_header_text = any(p.text.strip() for p in header.paragraphs)
            has_footer_text = any(p.text.strip() for p in footer.paragraphs)

            if "headers" in clear_rules:
                self._check("Force Clear: Headers",
                            not has_header_text,
                            "Headers empty (correct)",
                            "Header text still present")
            if "footers" in clear_rules:
                self._check("Force Clear: Footers",
                            not has_footer_text,
                            "Footers empty (correct)",
                            "Footer text still present")

    def _check_content_compliance(self):
        black = RGBColor(0, 0, 0)

        colored_text = 0
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb and \
                   run.font.color.rgb != black and str(run.font.color.rgb) != 'Auto':
                    colored_text += 1
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.color and run.font.color.rgb and \
                               run.font.color.rgb != black and str(run.font.color.rgb) != 'Auto':
                                colored_text += 1

        self._check("Content: No Colored Text",
                    colored_text == 0,
                    f"0 colored text runs (clean)",
                    f"{colored_text} colored text runs remain")

        colored_cells = 0
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill.upper() not in ('FFFFFF', 'AUTO', 'NONE', ''):
                                colored_cells += 1

        self._check("Content: No Colored Cell Backgrounds",
                    colored_cells == 0,
                    f"0 colored cell backgrounds (clean)",
                    f"{colored_cells} colored cell backgrounds remain")

        para_borders = 0
        for para in self.doc.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None and len(list(pBdr)) > 0:
                    para_borders += 1

        self._check("Content: No Paragraph Borders",
                    para_borders == 0,
                    f"0 paragraph borders (clean)",
                    f"{para_borders} paragraph borders remain")

    def _check_cover_page(self):
        """Check cover page compliance with industry standard layout rules."""
        cover_rules = self.template.get("cover_page_rules", {})
        if not cover_rules:
            return

        disallowed = cover_rules.get("disallowed_elements", [])

        first_heading_idx = None
        for i, para in enumerate(self.doc.paragraphs):
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name or "title" in style_name:
                first_heading_idx = i
                break

        if first_heading_idx is None:
            self._warn("Cover Page", "No heading found - cannot identify cover region")
            return

        cover_paras = []
        for i, para in enumerate(self.doc.paragraphs):
            if i >= first_heading_idx:
                break
            if para.text.strip():
                cover_paras.append((i, para))

        # Check navigation labels
        if "navigation_labels" in disallowed:
            nav_found = False
            for idx, para in cover_paras:
                text = para.text.strip()
                dot_chars = ['\u00b7', '\u2022', '\u2502', '\u2503', '\u2016']
                for dc in dot_chars:
                    if text.count(dc) >= 2:
                        nav_found = True
                        break
            # 内容保护：疑似导航行默认保留待确认 → 存在时为 WARN 而非 FAIL
            if nav_found:
                self._warn("Cover: Navigation Labels",
                           "疑似导航行已保留，请人工确认是否删除")
            else:
                self._check("Cover: No Navigation Labels", True,
                            "No navigation labels found (clean)", "")

        # Check data tables
        if "data_tables" in disallowed:
            first_heading_p = self.doc.paragraphs[first_heading_idx]._p
            body = self.doc.element.body
            body_children = list(body)
            try:
                heading_order = body_children.index(first_heading_p)
            except ValueError:
                heading_order = len(body_children)

            cover_tables = 0
            for table in self.doc.tables:
                tbl = table._element
                try:
                    if body_children.index(tbl) < heading_order:
                        cover_tables += 1
                except ValueError:
                    pass

            # 内容保护：封面表格默认保留待确认 → 存在时为 WARN 而非 FAIL
            if cover_tables > 0:
                self._warn("Cover: Data Tables",
                           f"{cover_tables} 个封面表格已保留，请人工确认是否删除")
            else:
                self._check("Cover: No Data Tables", True,
                            "No data tables on cover (clean)", "")

        # Check date position
        layout = cover_rules.get("layout", {})
        if layout and cover_paras:
            date_idx = None
            for idx, (para_idx, para) in enumerate(cover_paras):
                text = para.text.strip()
                if any(kw in text for kw in ['年', '月', '日']) and len(text) <= 20:
                    date_idx = idx
                    break

            if date_idx is not None and len(cover_paras) > 2:
                is_at_bottom = date_idx >= len(cover_paras) * 2 // 3
                self._check("Cover: Date at Bottom",
                            is_at_bottom,
                            f"Date at position {date_idx+1}/{len(cover_paras)} (bottom)",
                            f"Date at position {date_idx+1}/{len(cover_paras)} (not at bottom)")

            # Check spacing
            min_spacing = layout.get("min_empty_lines_between_title_and_date", 0)
            if min_spacing > 0 and date_idx is not None and len(cover_paras) >= 2:
                title_idx = cover_paras[0][0]
                date_para_idx = cover_paras[date_idx][0]
                empty_count = sum(1 for i in range(title_idx + 1, date_para_idx)
                                   if i < len(self.doc.paragraphs) and
                                   not self.doc.paragraphs[i].text.strip())
                self._check("Cover: Title-Date Spacing",
                            empty_count >= min_spacing,
                            f"{empty_count} empty lines between title and date (>= {min_spacing})",
                            f"Only {empty_count} empty lines (need {min_spacing})")

    def to_dict(self):
        """Stable machine-readable summary for CI/test pipelines."""
        return {
            "industry": self.template.get("industry_name", "Unknown"),
            "standard": self.template.get("standard", ""),
            "passed": self.passed,
            "failed": self.failed,
            "warnings": self.warnings,
            "ok": self.failed == 0,
            "checks": [
                {"status": status, "name": name, "detail": detail}
                for status, name, detail in self.results
            ],
        }

    def _report(self):
        print(f"\n{'='*60}")
        print(f"  Format Validation Report")
        print(f"  Industry: {self.template.get('industry_name', 'Unknown')}")
        print(f"  Standard: {self.template.get('standard', '')}")
        print(f"{'='*60}\n")

        for status, name, detail in self.results:
            icon = {"PASS": "[OK]", "FAIL": "[XX]", "WARN": "[!!]"}[status]
            print(f"  {icon} {name}: {detail}")

        print(f"\n{'─'*60}")
        print(f"  Total: {self.passed + self.failed + self.warnings}  "
              f"PASS: {self.passed}  FAIL: {self.failed}  WARN: {self.warnings}")

        if self.failed > 0:
            print(f"\n  RESULT: FAILED - {self.failed} check(s) failed")
        elif self.warnings > 0:
            print(f"\n  RESULT: PASSED (with warnings)")
        else:
            print(f"\n  RESULT: ALL CHECKS PASSED")


def main():
    parser = argparse.ArgumentParser(description="DOCX Format Validator")
    parser.add_argument("--input", required=True, help="Output docx file to validate")
    parser.add_argument("--template", required=True, help="Template JSON file path")
    parser.add_argument("--json", help="Also write machine-readable JSON summary to this path")
    args = parser.parse_args()

    validator = FormatValidator(args.input, args.template)
    validator.run()

    if args.json:
        with open(args.json, "w", encoding="utf-8") as f:
            json.dump(validator.to_dict(), f, ensure_ascii=False, indent=2)

    sys.exit(1 if validator.failed > 0 else 0)


if __name__ == "__main__":
    main()
