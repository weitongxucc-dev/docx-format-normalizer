#!/usr/bin/env python3
"""
DOCX Multi-Industry Document Format Normalizer v2
Processes .docx files to normalize formatting according to industry-specific standards.

Key improvements over v1:
- Document structure analysis (cover page / TOC / body detection)
- Style-level modification (styles.xml) in addition to run-level
- Cover page protection (preserve alignment, no body indent)
- Title heading support (title_heading field)
- Unified heading detection (Word styles + text patterns)
- Proper paragraph spacing clearance for fixed line spacing
- Table font inheritance from body text
- Complete header/footer handling
- Multiple regex patterns per hierarchy level

Usage:
    python3 docx_formatter.py --input input.docx --template templates/government_document.json --output output.docx
    python3 docx_formatter.py --config <config.json>
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Mm, Cm, Inches, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Run: bash scripts/setup.sh")
    print("Or: pip install 'python-docx>=0.8.11' 'lxml>=4.9,<6'")
    sys.exit(1)


class DocxFormatter:
    """Main processing engine for docx format normalization."""

    def __init__(self, config):
        self.config = config
        self.template = config.get("template", {})
        self.modifications = []
        self.doc = None
        self._para_roles = []
        self._title_para_idx = None

    @property
    def industry_name(self):
        return self.template.get("industry_name", "Unknown")

    def load_document(self, input_path):
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"输入文件不存在: {input_path}")
        if not input_path.lower().endswith('.docx'):
            raise ValueError(f"仅支持 .docx 格式文件，当前文件: {os.path.basename(input_path)}")
        try:
            import zipfile
            if not zipfile.is_zipfile(input_path):
                raise ValueError(f"文件不是有效的 docx 格式（ZIP 容器损坏）: {input_path}")
        except ImportError:
            pass
        try:
            self.doc = Document(input_path)
        except Exception as e:
            raise RuntimeError(f"解析 docx 文件失败: {e}")

    def run(self, input_path, output_path):
        self.load_document(input_path)
        self._analyze_document_structure()
        self._audit_and_cleanup()
        # 封面清理可能增删段落，必须重算角色，避免索引错位
        self._analyze_document_structure()
        self._normalize_list_paragraphs()
        self._format_cover_page()
        self._apply_styles()
        self._apply_page_setup()
        self._apply_title()
        self._apply_body_text()
        self._apply_headings()
        self._apply_hierarchy_fonts()
        self._apply_header_footer()
        self._apply_tables()
        self._apply_force_clear()
        self._apply_page_zones_if_needed()
        self.doc.save(output_path)
        self._cleanup_theme_in_output(output_path)
        return self.modifications

    def _cleanup_theme_in_output(self, output_path):
        """Post-save: clean up ALL color attributes in the docx package."""
        try:
            import zipfile
            import lxml.etree as etree
            import tempfile
            import os
            import shutil

            tmpdir = tempfile.mkdtemp(prefix="docx_color_")
            temp_out = os.path.join(tmpdir, "out.docx")

            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            safe_colors = {'000000', 'FFFFFF', 'AUTO', 'NONE', ''}

            theme_modified = False
            border_color_count = 0
            shd_count = 0

            with zipfile.ZipFile(output_path, 'r') as zin:
                with zipfile.ZipFile(temp_out, 'w', zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)

                        if not item.filename.endswith('.xml'):
                            zout.writestr(item, data)
                            continue

                        try:
                            root = etree.fromstring(data)
                        except Exception:
                            zout.writestr(item, data)
                            continue

                        changed = False

                        if 'theme' in item.filename:
                            for srgb in root.iter(f'{{{a_ns}}}srgbClr'):
                                val = srgb.get('val', '')
                                if val.upper() not in safe_colors:
                                    srgb.set('val', '000000')
                                    changed = True
                                    theme_modified = True

                        if item.filename in ('word/document.xml', 'word/styles.xml') or \
                           item.filename.startswith('word/header') or \
                           item.filename.startswith('word/footer'):
                            for elem in root.iter():
                                for attr_name in list(elem.attrib.keys()):
                                    short = attr_name.split('}')[-1] if '}' in attr_name else attr_name
                                    if short == 'color':
                                        val = elem.attrib[attr_name]
                                        if val.upper() not in safe_colors:
                                            elem.attrib[attr_name] = '000000'
                                            changed = True
                                            border_color_count += 1
                                    elif short == 'fill':
                                        val = elem.attrib[attr_name]
                                        if val.upper() not in safe_colors and \
                                           val.upper() not in ('BFBFBF', 'D9D9D9', 'F2F2F2'):
                                            elem.attrib[attr_name] = 'FFFFFF'
                                            changed = True
                                            shd_count += 1

                            for elem in root.iter():
                                for attr_name in list(elem.attrib.keys()):
                                    short = attr_name.split('}')[-1] if '}' in attr_name else attr_name
                                    if short == 'themeColor':
                                        val = elem.attrib[attr_name]
                                        if val not in ('text1', 'dark1', 'background1', 'light1'):
                                            elem.attrib[attr_name] = 'text1'
                                            changed = True

                            for rFonts in root.iter(f'{{{w_ns}}}rFonts'):
                                for attr_name in list(rFonts.attrib.keys()):
                                    short = attr_name.split('}')[-1] if '}' in attr_name else attr_name
                                    val = rFonts.attrib[attr_name]
                                    if val in ('Noto Sans CJK SC', 'Noto Sans SC', 'Noto Serif CJK SC'):
                                        rFonts.attrib[attr_name] = '宋体'
                                        changed = True
                                    elif val in ('MS ゴシック', 'MS ゴシック', 'MS Gothic'):
                                        rFonts.attrib[attr_name] = '宋体'
                                        changed = True
                                    elif val in ('MS 明朝', 'MS Mincho', 'MS 明朝'):
                                        rFonts.attrib[attr_name] = '宋体'
                                        changed = True
                                    elif val == 'Courier':
                                        rFonts.attrib[attr_name] = 'Times New Roman'
                                        changed = True

                        if changed:
                            zout.writestr(item, etree.tostring(
                                root, xml_declaration=True,
                                encoding='UTF-8', standalone=True))
                        else:
                            zout.writestr(item, data)

            shutil.copy2(temp_out, output_path)
            shutil.rmtree(tmpdir)

            if border_color_count:
                self._record("文档审查", "边框颜色",
                             f"{border_color_count}处彩色边框", "已全部改为黑色")
            if shd_count:
                self._record("文档审查", "样式阴影",
                             f"{shd_count}处彩色阴影", "已清除")
            if theme_modified:
                self._record("文档审查", "主题配色",
                             "彩色主题色(accent/hlink)", "已全部改为黑色")

        except Exception as e:
            self._record("文档审查", "深度颜色清理",
                         "尝试修改XML", f"错误: {e}")

    # ── Document Audit & Cleanup ──────────────────────────

    def _audit_and_cleanup(self):
        """Audit document for non-compliant content and auto-fix issues."""
        audit_issues = []

        colored_text = self._audit_colored_text()
        colored_cells = self._audit_colored_cells()
        nonstandard_sizes = self._audit_font_sizes()
        para_borders = self._audit_paragraph_borders()
        cover_issues = self._audit_cover_page()

        if colored_text:
            audit_issues.append({
                "issue": "colored_text",
                "count": colored_text,
                "detail": f"发现{colored_text}处彩色文字",
                "action": "已全部改为黑色"
            })
            self._cleanup_colored_text()

        if colored_cells:
            audit_issues.append({
                "issue": "colored_cell_background",
                "count": colored_cells,
                "detail": f"发现{colored_cells}个彩色背景单元格",
                "action": "已全部清除背景色"
            })
            self._cleanup_colored_cells()

        if para_borders:
            audit_issues.append({
                "issue": "paragraph_border",
                "count": para_borders,
                "detail": f"发现{para_borders}处段落装饰线",
                "action": "已全部删除"
            })
            self._cleanup_paragraph_borders()

        if cover_issues:
            audit_issues.append({
                "issue": "cover_page_nonstandard",
                "count": len(cover_issues),
                "detail": "封面存在非标准元素: " + "; ".join(cover_issues),
                "action": "已清理导航标签和装饰表格"
            })
            self._cleanup_cover_page()

        if nonstandard_sizes:
            audit_issues.append({
                "issue": "nonstandard_font_size",
                "count": nonstandard_sizes,
                "detail": f"发现{nonstandard_sizes}处非标准字号",
                "action": "将在格式化阶段统一"
            })

        if audit_issues:
            parts = []
            if colored_text: parts.append(f"彩色文字{colored_text}")
            if colored_cells: parts.append(f"彩色背景{colored_cells}")
            if para_borders: parts.append(f"段落装饰线{para_borders}")
            if cover_issues: parts.append(f"封面非标准{len(cover_issues)}项")
            if nonstandard_sizes: parts.append(f"非标准字号{nonstandard_sizes}")
            self._record("文档审查", f"发现{len(audit_issues)}类问题",
                         "/".join(parts),
                         "已自动清理，详见报告")

    def _audit_cover_page(self):
        """Audit cover page for non-standard elements and layout compliance."""
        issues = []
        cover_rules = self.template.get("cover_page_rules", {})
        if not cover_rules:
            return issues

        disallowed = cover_rules.get("disallowed_elements", [])

        cover_paras = []
        for i, para in enumerate(self.doc.paragraphs):
            role = self._para_roles[i] if i < len(self._para_roles) else "body"
            if role == "cover" and para.text.strip():
                cover_paras.append((i, para))

        # 1. Navigation labels
        if "navigation_labels" in disallowed:
            for idx, para in cover_paras:
                text = para.text.strip()
                dot_chars = ['·', '•', '│', '丨', '┃', '╎', '┆']
                for dc in dot_chars:
                    if text.count(dc) >= 2:
                        issues.append(f"导航标签: {text[:40]}")
                        break
                else:
                    if '|' in text and text.count('|') >= 2:
                        issues.append(f"导航标签: {text[:40]}")

        # 2. Data tables on cover (tables before first heading)
        if "data_tables" in disallowed:
            first_heading_idx = None
            for i, para in enumerate(self.doc.paragraphs):
                if self._is_heading_paragraph(para):
                    first_heading_idx = i
                    break

            if first_heading_idx is not None:
                first_heading_p = self.doc.paragraphs[first_heading_idx]._p
                body = self.doc.element.body
                body_children = list(body)
                try:
                    heading_order = body_children.index(first_heading_p)
                except ValueError:
                    heading_order = len(body_children)

                cover_table_count = 0
                for table in self.doc.tables:
                    tbl = table._element
                    try:
                        tbl_order = body_children.index(tbl)
                        if tbl_order < heading_order:
                            cover_table_count += 1
                    except ValueError:
                        pass

                if cover_table_count > 0:
                    issues.append(f"封面上有{cover_table_count}个数据表格")

        # 3. Text boxes / shapes on cover
        if "text_boxes" in disallowed:
            first_heading_idx = None
            for i, para in enumerate(self.doc.paragraphs):
                if self._is_heading_paragraph(para):
                    first_heading_idx = i
                    break

            if first_heading_idx is not None:
                first_heading_p = self.doc.paragraphs[first_heading_idx]._p
                body = self.doc.element.body
                body_children = list(body)
                try:
                    heading_order = body_children.index(first_heading_p)
                except ValueError:
                    heading_order = len(body_children)

                textbox_count = 0
                for child in body_children[:heading_order]:
                    drawings = child.findall('.//' + qn('w:drawing'))
                    picts = child.findall('.//' + qn('w:pict'))
                    textbox_count += len(drawings) + len(picts)

                if textbox_count > 0:
                    issues.append(f"封面上有{textbox_count}个文本框/图形")

        # 4. Layout validation
        layout = cover_rules.get("layout", {})
        if layout and cover_paras:
            layout_issues = self._validate_cover_layout(cover_paras, layout)
            issues.extend(layout_issues)

        return issues

    def _validate_cover_layout(self, cover_paras, layout):
        """Validate cover page element positions against standard layout."""
        issues = []
        total = len(cover_paras)
        if total == 0:
            return issues

        # Check date position - should be at bottom
        date_keywords = ['年', '月', '日']
        date_idx = None
        for idx, (para_idx, para) in enumerate(cover_paras):
            text = para.text.strip()
            if any(kw in text for kw in date_keywords):
                if len(text) <= 20:
                    date_idx = idx
                    break

        if date_idx is not None and total > 2:
            if date_idx < total * 2 // 3:
                issues.append(
                    f"日期位置不在底部（当前第{date_idx+1}/{total}段，应在底部）")

        # Check spacing between title and date
        min_spacing = layout.get("min_empty_lines_between_title_and_date", 0)
        if min_spacing > 0 and date_idx is not None and total >= 2:
            title_para_idx = cover_paras[0][0]
            date_para_idx = cover_paras[date_idx][0]
            empty_count = 0
            for i in range(title_para_idx + 1, date_para_idx):
                if i < len(self._para_roles):
                    if self._para_roles[i] == "empty":
                        empty_count += 1
            if empty_count < min_spacing:
                issues.append(
                    f"标题与日期间留白不足（当前{empty_count}行，需{min_spacing}行）")

        return issues

    def _cleanup_cover_page(self):
        """Remove non-standard cover page elements and restructure layout."""
        cover_rules = self.template.get("cover_page_rules", {})
        if not cover_rules:
            return

        disallowed = cover_rules.get("disallowed_elements", [])

        # 1. Remove navigation labels
        if "navigation_labels" in disallowed:
            paras_to_remove = []
            for i, para in enumerate(self.doc.paragraphs):
                role = self._para_roles[i] if i < len(self._para_roles) else "body"
                if role != "cover":
                    continue
                text = para.text.strip()
                should_remove = False
                dot_chars = ['·', '•', '│', '丨', '┃', '╎', '┆']
                for dc in dot_chars:
                    if text.count(dc) >= 2:
                        should_remove = True
                        break
                if not should_remove and '|' in text and text.count('|') >= 2:
                    should_remove = True

                if should_remove:
                    paras_to_remove.append((para, text))

            for para, text in paras_to_remove:
                p_element = para._element
                p_element.getparent().remove(p_element)
                self._record("文档审查", "封面清理",
                             f"删除导航标签: {text[:30]}", "已删除")

        # 2. Remove data tables from cover
        if "data_tables" in disallowed:
            first_heading_idx = None
            for i, para in enumerate(self.doc.paragraphs):
                if self._is_heading_paragraph(para):
                    first_heading_idx = i
                    break

            if first_heading_idx is not None:
                first_heading_p = self.doc.paragraphs[first_heading_idx]._p
                body = self.doc.element.body
                body_children = list(body)
                try:
                    heading_order = body_children.index(first_heading_p)
                except ValueError:
                    heading_order = len(body_children)

                tables_to_remove = []
                for table in self.doc.tables:
                    tbl = table._element
                    try:
                        tbl_order = body_children.index(tbl)
                        if tbl_order < heading_order:
                            tables_to_remove.append(tbl)
                    except ValueError:
                        pass

                for tbl in tables_to_remove:
                    tbl.getparent().remove(tbl)
                    self._record("文档审查", "封面清理",
                                 "删除封面数据表格", "已删除")

        # 3. Remove text boxes / shapes from cover
        if "text_boxes" in disallowed:
            first_heading_idx = None
            for i, para in enumerate(self.doc.paragraphs):
                if self._is_heading_paragraph(para):
                    first_heading_idx = i
                    break

            if first_heading_idx is not None:
                first_heading_p = self.doc.paragraphs[first_heading_idx]._p
                body = self.doc.element.body
                body_children = list(body)
                try:
                    heading_order = body_children.index(first_heading_p)
                except ValueError:
                    heading_order = len(body_children)

                textbox_removed = 0
                for child in body_children[:heading_order]:
                    drawings = child.findall('.//' + qn('w:drawing'))
                    picts = child.findall('.//' + qn('w:pict'))
                    for elem in drawings + picts:
                        elem.getparent().remove(elem)
                        textbox_removed += 1

                if textbox_removed > 0:
                    self._record("文档审查", "封面清理",
                                 f"删除{textbox_removed}个文本框/图形", "已删除")

        # 4. Restructure cover layout (add spacing if needed)
        layout = cover_rules.get("layout", {})
        if layout:
            self._restructure_cover_layout(layout)

    def _restructure_cover_layout(self, layout):
        """Ensure proper spacing between cover elements."""
        cover_paras = []
        for i, para in enumerate(self.doc.paragraphs):
            role = self._para_roles[i] if i < len(self._para_roles) else "body"
            if role == "cover" and para.text.strip():
                cover_paras.append((i, para))

        if len(cover_paras) < 2:
            return

        min_spacing = layout.get("min_empty_lines_between_title_and_date", 0)
        if min_spacing <= 0:
            return

        title_idx = cover_paras[0][0]
        last_idx = cover_paras[-1][0]
        empty_count = 0
        for i in range(title_idx + 1, last_idx):
            if i < len(self._para_roles):
                if self._para_roles[i] == "empty":
                    empty_count += 1

        if empty_count >= min_spacing:
            return

        needed = min_spacing - empty_count
        last_para = cover_paras[-1][1]
        for _ in range(needed):
            new_p = OxmlElement('w:p')
            last_para._element.addprevious(new_p)
        self._record("文档审查", "封面布局调整",
                     f"标题与日期间补充{needed}个空行",
                     f"已补充至{min_spacing}行留白")

    def _audit_paragraph_borders(self):
        """Count paragraphs with decorative borders (lines below headings etc.)."""
        count = 0
        for para in self.doc.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None and len(list(pBdr)) > 0:
                    count += 1
        return count

    def _cleanup_paragraph_borders(self):
        """Remove all paragraph decorative borders."""
        for para in self.doc.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    pPr.remove(pBdr)

    def _audit_colored_text(self):
        """Count runs with non-black font color."""
        count = 0
        black = RGBColor(0, 0, 0)
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb and \
                   run.font.color.rgb != black and str(run.font.color.rgb) != 'Auto':
                    count += 1
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.color and run.font.color.rgb and \
                               run.font.color.rgb != black and str(run.font.color.rgb) != 'Auto':
                                count += 1
        return count

    def _audit_colored_cells(self):
        """Count table cells with non-white background shading."""
        count = 0
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill.upper() not in ('FFFFFF', 'AUTO', 'NONE', ''):
                                count += 1
        return count

    def _audit_font_sizes(self):
        """Count runs with font sizes not matching template body text size."""
        body = self.template.get("body_text", {})
        expected_size = body.get("font_size_pt")
        if not expected_size:
            return 0
        count = 0
        for para in self.doc.paragraphs:
            role = "body"
            for i, p in enumerate(self.doc.paragraphs):
                if p._p is para._p:
                    role = self._para_roles[i] if i < len(self._para_roles) else "body"
                    break
            if role == "cover":
                continue
            for run in para.runs:
                if run.font.size and abs(run.font.size.pt - expected_size) > 1:
                    count += 1
        return count

    def _cleanup_colored_text(self):
        """Set all non-black text to black, including style-level colors."""
        black = RGBColor(0, 0, 0)

        for style in self.doc.styles:
            try:
                if hasattr(style, 'font') and style.font.color and \
                   style.font.color.rgb and style.font.color.rgb != black and \
                   str(style.font.color.rgb) != 'Auto':
                    style.font.color.rgb = black
            except Exception:
                pass

        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb and \
                   run.font.color.rgb != black and str(run.font.color.rgb) != 'Auto':
                    run.font.color.rgb = black
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.font.color and run.font.color.rgb and \
                               run.font.color.rgb != black and str(run.font.color.rgb) != 'Auto':
                                run.font.color.rgb = black

        rPr_template = qn('w:rPr')
        color_tag = qn('w:color')
        for style_elem in self.doc.styles.element.iterchildren():
            rPr = style_elem.find(rPr_template)
            if rPr is not None:
                color = rPr.find(color_tag)
                if color is not None:
                    val = color.get(qn('w:val'))
                    if val and val.upper() not in ('000000', 'AUTO', 'NONE', ''):
                        color.set(qn('w:val'), '000000')

    def _cleanup_colored_cells(self):
        """Clear all non-white cell background shading."""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill.upper() not in ('FFFFFF', 'AUTO', 'NONE', ''):
                                shd.set(qn('w:fill'), 'FFFFFF')
                                shd.set(qn('w:val'), 'clear')

    # ── Document Structure Analysis ────────────────────────

    def _analyze_document_structure(self):
        """Classify each paragraph as cover, toc, or body."""
        first_heading_idx = None
        for i, para in enumerate(self.doc.paragraphs):
            if self._is_heading_paragraph(para):
                first_heading_idx = i
                break

        self._para_roles = []
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            style_name = (para.style.name or "").lower() if para.style else ""

            if not text:
                self._para_roles.append("empty")
            elif "toc" in style_name:
                self._para_roles.append("toc")
            elif first_heading_idx is not None and i < first_heading_idx:
                self._para_roles.append("cover")
            else:
                self._para_roles.append("body")

    def _get_role(self, para):
        for i, p in enumerate(self.doc.paragraphs):
            if p._p is para._p:
                return self._para_roles[i] if i < len(self._para_roles) else "body"
        return "body"

    # ── Style-Level Modification ───────────────────────────

    def _apply_styles(self):
        """Modify style definitions in styles.xml for persistent formatting."""
        body = self.template.get("body_text")
        if body:
            self._update_style("Normal", body)

        title_heading = self.template.get("title_heading")
        if title_heading:
            self._update_style("Title", title_heading)
        else:
            headings = self.template.get("headings", [])
            for h in headings:
                if h.get("level", 0) == 0:
                    self._update_style("Title", h)
                    break

        headings = self.template.get("headings", [])
        for h in headings:
            level = h.get("level", 0)
            if level > 0:
                style_name = f"Heading {level}"
                self._update_style(style_name, h)

    def _update_style(self, style_name, fmt):
        """Update a paragraph style definition with formatting parameters."""
        try:
            style = self.doc.styles[style_name]
        except KeyError:
            return

        font_cn = fmt.get("font_cn", fmt.get("font", "宋体"))
        font_en = fmt.get("font_en", "Times New Roman")
        font_size = fmt.get("font_size_pt")
        text_color = fmt.get("text_color", "000000")

        if font_cn:
            self._set_style_east_asia_font(style, font_cn, font_en)

        if font_size:
            style.font.size = Pt(font_size)

        if font_en:
            style.font.name = font_en

        if text_color:
            try:
                style.font.color.rgb = RGBColor.from_string(text_color)
            except Exception:
                pass

        bold = fmt.get("bold")
        if bold is not None:
            style.font.bold = bold

        alignment = fmt.get("alignment")
        if alignment:
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            if alignment in align_map:
                style.paragraph_format.alignment = align_map[alignment]

        pf = style.paragraph_format
        line_rule = fmt.get("line_spacing_rule")
        if line_rule == "1.5x":
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif line_rule == "single":
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif line_rule == "fixed":
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(fmt.get("line_spacing_exact_pt", 28.8))

        if line_rule == "fixed":
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
        else:
            sb = fmt.get("space_before_lines")
            sa = fmt.get("space_after_lines")
            base_size = font_size or 14
            if sb is not None:
                pf.space_before = Pt(base_size * sb)
            if sa is not None:
                pf.space_after = Pt(base_size * sa)

        first_indent = fmt.get("first_line_indent_chars", 0)
        if first_indent and first_indent > 0 and line_rule != "fixed":
            base_size = font_size or 14
            pf.first_line_indent = Pt(base_size * first_indent)

        self._record("样式", f"{style_name} 样式定义", "—",
                     f"{font_cn} {font_size}pt {line_rule or ''}")

    def _set_style_east_asia_font(self, style, font_cn, font_en):
        """Set eastAsia font on a style element, with macOS fallback."""
        rpr = style.element.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            style.element.append(rpr)
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_cn)
        rFonts.set(qn('w:ascii'), font_en or font_cn)
        rFonts.set(qn('w:hAnsi'), font_en or font_cn)
        font_fallback = self.template.get("font_fallback", {})
        alt_font = font_fallback.get(font_cn)
        if alt_font:
            rFonts.set(qn('w:cs'), alt_font)

    # ── Page Setup ──────────────────────────────────────────

    def _apply_page_setup(self):
        page_setup = self.template.get("page_setup")
        if not page_setup:
            return

        for section in self.doc.sections:
            if page_setup.get("paper_size") == "A4":
                section.page_width = Mm(210)
                section.page_height = Mm(297)
                self._record("页面设置", "纸张大小", "—", "A4 (210×297mm)")

            if page_setup.get("orientation") == "portrait":
                section.orientation = WD_ORIENT.PORTRAIT

            margins = page_setup.get("margins", {})
            margin_map = {
                "top_mm": ("top_margin", Mm),
                "bottom_mm": ("bottom_margin", Mm),
                "left_mm": ("left_margin", Mm),
                "right_mm": ("right_margin", Mm),
            }
            for key, (attr, unit) in margin_map.items():
                if key in margins:
                    new_val = margins[key]
                    old_val = getattr(section, attr)
                    old_mm = round(old_val / 36000, 1) if old_val else 0
                    if abs(old_mm - new_val) > 0.5:
                        setattr(section, attr, unit(new_val))
                        self._record("页面设置", f"边距({key})",
                                     f"{old_mm}mm", f"{new_val}mm")

            lines_per_page = page_setup.get("lines_per_page")
            chars_per_line = page_setup.get("chars_per_line")
            if lines_per_page and chars_per_line:
                self._set_grid_lines(section, lines_per_page, chars_per_line)

    def _set_grid_lines(self, section, lines_per_page, chars_per_line):
        sectPr = section._sectPr
        docGrid = sectPr.find(qn('w:docGrid'))
        if docGrid is None:
            docGrid = OxmlElement('w:docGrid')
            sectPr.append(docGrid)
        docGrid.set(qn('w:type'), 'lines')
        docGrid.set(qn('w:linePitch'), '312')
        self._record("页面设置", "文档网格",
                     "—", f"{lines_per_page}行×{chars_per_line}字/页")

    # ── Title Heading ────────────────────────────────────────

    def _apply_title(self):
        """Apply title_heading configuration to the document title."""
        title_cfg = self.template.get("title_heading")
        if not title_cfg:
            return

        title_para = None
        title_idx = None
        for i, para in enumerate(self.doc.paragraphs):
            role = self._para_roles[i] if i < len(self._para_roles) else "body"
            if role == "cover" and para.text.strip():
                style_name = (para.style.name or "").lower() if para.style else ""
                if "title" in style_name:
                    title_para = para
                    title_idx = i
                    break

        if title_para is None:
            for i, para in enumerate(self.doc.paragraphs):
                if para.text.strip():
                    align = para.alignment
                    has_large_font = False
                    for run in para.runs:
                        if run.font.size and run.font.size > Pt(16):
                            has_large_font = True
                            break
                    if align == WD_ALIGN_PARAGRAPH.CENTER or has_large_font or \
                       (para.style and "title" in (para.style.name or "").lower()):
                        title_para = para
                        title_idx = i
                        break

        if title_para is None:
            return

        self._title_para_idx = title_idx

        font_cn = title_cfg.get("font_cn", title_cfg.get("font", "黑体"))
        font_en = title_cfg.get("font_en", "Times New Roman")
        font_size = title_cfg.get("font_size_pt", 22)
        bold = title_cfg.get("bold", False)
        alignment = title_cfg.get("alignment", "center")

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        if alignment in align_map:
            title_para.alignment = align_map[alignment]

        title_para.paragraph_format.first_line_indent = None

        for run in title_para.runs:
            self._format_run(run, font_cn, font_en, font_size, "000000")
            run.font.bold = bold

        preview = title_para.text[:30] + ("..." if len(title_para.text) > 30 else "")
        self._record("文档标题", preview, "—",
                     f"{font_cn} {font_size}pt {'加粗' if bold else ''} {alignment}")

    # ── Body Text ───────────────────────────────────────────

    def _normalize_list_paragraphs(self):
        """Remove non-standard list bullets (•/▪/数字编号) by converting
        List* styles to Normal, then they are formatted as standard body text."""
        count = 0
        for para in self.doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""
            if "list" not in style_name:
                continue
            try:
                para.style = self.doc.styles['Normal']
            except KeyError:
                pass
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)
            count += 1

        if count:
            self._record("列表规范化", f"{count}个列表段落",
                         "项目符号•/编号", "转为标准正文(去除非标符号)")

    def _apply_body_text(self):
        body = self.template.get("body_text")
        if not body:
            return

        for i, para in enumerate(self.doc.paragraphs):
            if self._is_heading_paragraph(para):
                continue

            if i == self._title_para_idx:
                continue

            role = self._para_roles[i] if i < len(self._para_roles) else "body"

            if role == "empty":
                continue
            elif role == "cover":
                continue  # 封面由 _format_cover_page 专门排版
            elif role == "toc":
                continue
            else:
                self._format_paragraph(para, body)

    def _format_cover_page(self):
        """Format the cover page to a standard, professional layout.

        Classifies each cover paragraph into an element type (title /
        subtitle / meta / date), then applies the template's cover element
        spec (font, size, alignment) plus vertical spacing so the cover is
        centered and visually distributed instead of left-aligned and crammed.
        """
        cover_rules = self.template.get("cover_page_rules", {})
        layout = cover_rules.get("layout", {})
        elements = {e.get("name"): e for e in layout.get("elements", [])}
        distribution = layout.get("vertical_distribution", "balanced")
        body = self.template.get("body_text", {})

        cover_paras = []
        for i, para in enumerate(self.doc.paragraphs):
            role = self._para_roles[i] if i < len(self._para_roles) else "body"
            if role == "cover" and para.text.strip():
                cover_paras.append(para)

        if not cover_paras:
            return

        import re as _re
        date_re = _re.compile(r'(\d{4}[-年/.]\d{1,2}[-月/.]\d{1,2}|\d{4}年|\d{1,2}月)')

        n = len(cover_paras)
        for idx, para in enumerate(cover_paras):
            if idx == 0:
                etype = "title"
            elif idx == n - 1 and date_re.search(para.text):
                etype = "date"
            elif idx == 1:
                etype = "subtitle"
            else:
                etype = "meta"

            spec = elements.get(etype, {})
            if etype == "title":
                font_cn = spec.get("font", "小标宋体")
                size = spec.get("size_pt", 22)
                align = spec.get("alignment", "center")
                sb, sa = 72, 24
            elif etype == "subtitle":
                font_cn = spec.get("font", "黑体")
                size = spec.get("size_pt", 16)
                align = spec.get("alignment", "center")
                sb, sa = 12, 12
            elif etype == "date":
                font_cn = spec.get("font", body.get("font_cn", "仿宋_GB2312"))
                size = spec.get("size_pt", 14)
                align = spec.get("alignment", "center")
                sb, sa = (36 if distribution == "balanced" else 12), 0
            else:
                font_cn = spec.get("font", body.get("font_cn", "仿宋_GB2312"))
                size = spec.get("size_pt", 14)
                align = spec.get("alignment", "center")
                sb, sa = 6, 0

            for run in para.runs:
                self._format_run(run, font_cn, "Times New Roman", size, "000000")

            align_map = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)

            pf = para.paragraph_format
            pf.first_line_indent = None
            pf.left_indent = None
            pf.space_before = Pt(sb)
            pf.space_after = Pt(sa)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

            preview = para.text[:25] + ("..." if len(para.text) > 25 else "")
            self._record("封面排版", preview, "左对齐/小字",
                         f"{etype}→{font_cn} {size}pt 居中")

    def _format_paragraph(self, para, fmt):
        font_cn = fmt.get("font_cn", "宋体")
        font_en = fmt.get("font_en", "Times New Roman")
        font_size = fmt.get("font_size_pt")
        line_spacing_rule = fmt.get("line_spacing_rule", "1.5x")
        first_indent = fmt.get("first_line_indent_chars", 0)
        alignment = fmt.get("alignment", "left")
        text_color = fmt.get("text_color", "000000")

        changes = []

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if alignment in align_map:
            old_align = para.alignment
            new_align = align_map[alignment]
            if old_align != new_align:
                para.alignment = new_align
                changes.append(f"对齐→{alignment}")

        pf = para.paragraph_format
        if line_spacing_rule == "1.5x":
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            changes.append("行距→1.5倍")
        elif line_spacing_rule == "single":
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            changes.append("行距→单倍")
        elif line_spacing_rule == "fixed":
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            exact_pt = fmt.get("line_spacing_exact_pt", 28.8)
            pf.line_spacing = Pt(exact_pt)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            changes.append(f"行距→固定{exact_pt}pt 段间距→0")

        if first_indent and first_indent > 0:
            pf.first_line_indent = Pt(font_size * first_indent)
            changes.append(f"首行缩进→{first_indent}字符")
        else:
            pf.first_line_indent = None

        if line_spacing_rule != "fixed":
            space_before = fmt.get("space_before_lines")
            space_after = fmt.get("space_after_lines")
            if space_before is not None:
                pf.space_before = Pt(font_size * space_before)
            if space_after is not None:
                pf.space_after = Pt(font_size * space_after)

        for run in para.runs:
            run_changes = self._format_run(run, font_cn, font_en, font_size, text_color, fmt)
            changes.extend(run_changes)

        if changes:
            preview = para.text[:30] + ("..." if len(para.text) > 30 else "")
            self._record("正文段落", preview, "—", "; ".join(changes[:4]))

    def _format_run(self, run, font_cn, font_en, font_size, text_color, fmt=None):
        changes = []

        if font_cn:
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            old_ea = rFonts.get(qn('w:eastAsia'))
            if old_ea != font_cn:
                rFonts.set(qn('w:eastAsia'), font_cn)
                rFonts.set(qn('w:ascii'), font_en or font_cn)
                rFonts.set(qn('w:hAnsi'), font_en or font_cn)
                font_fallback = self.template.get("font_fallback", {})
                alt_font = font_fallback.get(font_cn)
                if alt_font:
                    rFonts.set(qn('w:cs'), alt_font)
                changes.append(f"字体→{font_cn}")

        if font_size and (run.font.size is None or run.font.size != Pt(font_size)):
            run.font.size = Pt(font_size)
            changes.append(f"字号→{font_size}pt")

        if text_color:
            color = RGBColor.from_string(text_color)
            if run.font.color is None or run.font.color.rgb is None or \
               run.font.color.rgb != color:
                run.font.color.rgb = color
                changes.append("颜色→黑色")

        if fmt and fmt.get("char_spacing_pt"):
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            spacing = rPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                rPr.append(spacing)
            spacing.set(qn('w:val'), str(int(fmt["char_spacing_pt"] * 20)))

        return changes

    # ── Headings ─────────────────────────────────────────────

    def _is_heading_paragraph(self, para):
        style_name = (para.style.name or "").lower() if para.style else ""
        return "heading" in style_name or "title" in style_name

    def _apply_headings(self):
        headings = self.template.get("headings", [])
        if not headings:
            return

        for para in self.doc.paragraphs:
            for h in headings:
                if self._match_heading(para, h):
                    self._format_heading(para, h)
                    break

    def _match_heading(self, para, h):
        style_name = (para.style.name or "").lower() if para.style else ""
        level = h.get("level", 0)
        if level == 0:
            return "title" in style_name
        return f"heading {level}" in style_name

    def _format_heading(self, para, h):
        font_cn = h.get("font_cn", h.get("font", "宋体"))
        font_en = h.get("font_en", "Times New Roman")
        font_size = h.get("font_size_pt", 14)
        bold = h.get("bold", False)
        alignment = h.get("alignment", "center")

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        if alignment in align_map:
            para.alignment = align_map[alignment]

        pf = para.paragraph_format
        pf.first_line_indent = None

        line_rule = h.get("line_spacing")
        if line_rule == "1.5x":
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif line_rule == "single":
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

        if h.get("space_before_lines") is not None:
            pf.space_before = Pt(font_size * h["space_before_lines"])
        if h.get("space_after_lines") is not None:
            pf.space_after = Pt(font_size * h["space_after_lines"])

        for run in para.runs:
            self._format_run(run, font_cn, font_en, font_size, "000000")
            run.font.bold = bold

        preview = para.text[:30] + ("..." if len(para.text) > 30 else "")
        self._record("标题", preview, "—",
                     f"{font_cn} {font_size}pt {'加粗' if bold else ''} {alignment}")

    # ── Hierarchy Fonts ─────────────────────────────────────

    def _apply_hierarchy_fonts(self):
        hier_fonts = self.template.get("hierarchy_fonts", [])
        if not hier_fonts:
            return

        for i, para in enumerate(self.doc.paragraphs):
            if self._is_heading_paragraph(para):
                continue

            role = self._para_roles[i] if i < len(self._para_roles) else "body"
            if role in ("cover", "toc", "empty"):
                continue

            text = para.text.strip()
            if not text:
                continue

            for h in hier_fonts:
                patterns = h.get("prefix_patterns", [])
                single_pattern = h.get("prefix_pattern", "")
                if single_pattern:
                    patterns = [single_pattern] + patterns

                matched = False
                for pat in patterns:
                    if pat and re.search(pat, text):
                        matched = True
                        break

                if matched:
                    font_cn = h.get("font", h.get("font_cn", "仿宋"))
                    font_en = h.get("font_en", "Times New Roman")
                    font_size = h.get("font_size_pt", 16)
                    bold = h.get("bold", False)

                    for run in para.runs:
                        self._format_run(run, font_cn, font_en, font_size, "000000")
                        run.font.bold = bold

                    pf = para.paragraph_format
                    pf.first_line_indent = None

                    preview = text[:30] + ("..." if len(text) > 30 else "")
                    self._record("层次序号", preview, "—",
                                 f"{font_cn} {font_size}pt")
                    break

    # ── Header / Footer ──────────────────────────────────────

    def _apply_header_footer(self):
        hf = self.template.get("header_footer")
        if not hf:
            return

        for section in self.doc.sections:
            if not hf.get("has_header", True):
                self._clear_section_header(section)
                self._record("页眉", "清除", "有内容", "已清空")

            if not hf.get("has_footer", True):
                self._clear_section_footer(section)
                self._record("页脚", "清除", "有内容", "已清空")

            if not hf.get("has_page_number", True):
                self._remove_page_numbers(section)
                self._record("页码", "清除", "有页码", "已移除")

            if hf.get("footer_type") == "page_number" and hf.get("has_page_number", True):
                self._set_page_number_format(section, hf)

            if hf.get("has_header", True):
                body = self.template.get("body_text", {})
                header_font = body.get("font_cn", "宋体")
                header_font_en = body.get("font_en", "Times New Roman")
                self._update_header_font(section, header_font, header_font_en)

    def _clear_section_header(self, section):
        header = section.header
        for para in header.paragraphs:
            for run in para.runs:
                run.text = ""
            for child in list(para._element):
                if child.tag == qn('w:r'):
                    para._element.remove(child)

    def _clear_section_footer(self, section):
        footer = section.footer
        for para in footer.paragraphs:
            for run in para.runs:
                run.text = ""
            for child in list(para._element):
                if child.tag == qn('w:r'):
                    para._element.remove(child)

    def _update_header_font(self, section, font_cn, font_en):
        header = section.header
        for para in header.paragraphs:
            for run in para.runs:
                r = run._element
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), font_cn)
                rFonts.set(qn('w:ascii'), font_en)
                rFonts.set(qn('w:hAnsi'), font_en)

    def _remove_page_numbers(self, section):
        for part in [section.header, section.footer]:
            for para in part.paragraphs:
                for run in para.runs:
                    for child in list(run._element):
                        if child.tag in (qn('w:fldSimple'), qn('w:fldChar'),
                                         qn('w:instrText')):
                            run._element.remove(child)

    def _set_page_number_format(self, section, hf):
        font_name = hf.get("page_number_font", "宋体")
        font_size = hf.get("page_number_size_pt", 14)
        odd_align = hf.get("page_number_odd_align", "right")

        footer = section.footer
        footer.is_linked_to_previous = False

        for para in footer.paragraphs:
            for child in list(para._element):
                if child.tag == qn('w:r'):
                    para._element.remove(child)

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        align_map = {
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
        }
        if odd_align in align_map:
            para.alignment = align_map[odd_align]

        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run.font.name = font_name
        run.font.size = Pt(font_size)

        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)

        self._record("页码", "设置页码格式", "—",
                     f"{font_name} {font_size}pt {odd_align}对齐")

    # ── Tables ───────────────────────────────────────────────

    def _apply_tables(self):
        table_cfg = self.template.get("tables")
        if not table_cfg:
            return

        body = self.template.get("body_text", {})
        default_font_cn = body.get("font_cn", "宋体")
        default_font_en = body.get("font_en", "Times New Roman")

        cell_font_cn = table_cfg.get("cell_font_cn", default_font_cn)
        cell_font_en = table_cfg.get("cell_font_en", default_font_en)
        cell_font_size = table_cfg.get("cell_font_size_pt")

        for table in self.doc.tables:
            border_width = table_cfg.get("border_width_pt", 0.5)
            self._set_table_borders(table, border_width)

            if table_cfg.get("table_alignment") == "center":
                table.alignment = 1

            self._fit_table_to_page(table)

            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if cell_font_cn:
                                self._format_run(run, cell_font_cn, cell_font_en,
                                                 cell_font_size, "000000")

        self._record("表格", f"共{len(self.doc.tables)}个表格", "—",
                     f"边框{table_cfg.get('border_width_pt', 0.5)}磅 "
                     f"字体→{cell_font_cn}")

    def _set_table_borders(self, table, width_pt):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        borders = tblPr.find(qn('w:tblBorders'))
        if borders is None:
            borders = OxmlElement('w:tblBorders')
            tblPr.append(borders)
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            elem = borders.find(qn(f'w:{edge}'))
            if elem is None:
                elem = OxmlElement(f'w:{edge}')
                borders.append(elem)
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), str(int(width_pt * 8)))
            elem.set(qn('w:color'), '000000')

    def _fit_table_to_page(self, table):
        """Scale table column widths so the whole table fits within the page
        text area (版心), preventing horizontal cutoff on a single page."""
        page_setup = self.template.get("page_setup", {})
        margins = page_setup.get("margins", {})
        left_mm = margins.get("left_mm", 28)
        right_mm = margins.get("right_mm", 26)
        page_w_mm = 210 if page_setup.get("paper_size", "A4") == "A4" else 210
        text_width_twips = int((page_w_mm - left_mm - right_mm) * 56.7)

        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(text_width_twips))

        layout = tblPr.find(qn('w:tblLayout'))
        if layout is None:
            layout = OxmlElement('w:tblLayout')
            tblPr.append(layout)
        layout.set(qn('w:type'), 'fixed')

        grid = tbl.find(qn('w:tblGrid'))
        if grid is not None:
            cols = grid.findall(qn('w:gridCol'))
            if cols:
                n = len(cols)
                widths = []
                for c in cols:
                    w = c.get(qn('w:w'))
                    widths.append(int(w) if w else text_width_twips // n)
                total = sum(widths) or 1
                new_widths = [max(int(w * text_width_twips / total), 400)
                              for w in widths]
                diff = text_width_twips - sum(new_widths)
                new_widths[-1] += diff
                for c, w in zip(cols, new_widths):
                    c.set(qn('w:w'), str(w))

        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        tcW.set(qn('w:type'), 'auto')
                        tcW.set(qn('w:w'), '0')

    # ── Force Clear (Dark-bid mode) ─────────────────────────

    def _apply_force_clear(self):
        clear_rules = self.template.get("force_clear", [])
        if not clear_rules:
            return

        for rule in clear_rules:
            if rule == "bold":
                self._clear_bold()
            elif rule == "italic":
                self._clear_italic()
            elif rule == "underline":
                self._clear_underline()
            elif rule == "strikethrough":
                self._clear_strikethrough()
            elif rule in ("colored_text", "colored_background"):
                self._clear_colors()
            elif rule == "headers":
                self._clear_headers()
            elif rule == "footers":
                self._clear_footers()
            elif rule == "page_numbers":
                self._clear_page_numbers()
            elif rule == "table_of_contents":
                self._clear_toc()
            elif rule == "hyperlink_coloring":
                self._clear_hyperlink_colors()
            elif rule == "custom_char_spacing":
                self._clear_char_spacing()
            elif rule == "custom_char_scale":
                self._clear_char_scale()

    def _iter_all_runs(self):
        for para in self.doc.paragraphs:
            for run in para.runs:
                yield run
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            yield run

    def _clear_bold(self):
        count = sum(1 for run in self._iter_all_runs() if run.font.bold)
        for run in self._iter_all_runs():
            if run.font.bold:
                run.font.bold = False
        for style in self.doc.styles:
            try:
                if style.font.bold:
                    style.font.bold = False
            except Exception:
                pass
        if count:
            self._record("强制清除", "加粗", f"{count}处", "全部清除")

    def _clear_italic(self):
        count = sum(1 for run in self._iter_all_runs() if run.font.italic)
        if count:
            for run in self._iter_all_runs():
                if run.font.italic:
                    run.font.italic = False
            self._record("强制清除", "倾斜", f"{count}处", "全部清除")

    def _clear_underline(self):
        from docx.enum.text import WD_UNDERLINE
        count = sum(1 for run in self._iter_all_runs()
                     if run.font.underline and run.font.underline != WD_UNDERLINE.NONE)
        if count:
            for run in self._iter_all_runs():
                if run.font.underline and run.font.underline != WD_UNDERLINE.NONE:
                    run.font.underline = WD_UNDERLINE.NONE
            self._record("强制清除", "下划线", f"{count}处", "全部清除")

    def _clear_strikethrough(self):
        count = sum(1 for run in self._iter_all_runs() if run.font.strike)
        if count:
            for run in self._iter_all_runs():
                if run.font.strike:
                    run.font.strike = False
            self._record("强制清除", "删除线", f"{count}处", "全部清除")

    def _clear_colors(self):
        black = RGBColor(0, 0, 0)
        count = sum(1 for run in self._iter_all_runs()
                     if run.font.color and run.font.color.rgb and
                     run.font.color.rgb != black)
        if count:
            for run in self._iter_all_runs():
                if run.font.color and run.font.color.rgb and run.font.color.rgb != black:
                    run.font.color.rgb = black
            self._record("强制清除", "彩色字体", f"{count}处", "全部改为黑色")

    def _clear_headers(self):
        for section in self.doc.sections:
            self._clear_section_header(section)
        self._record("强制清除", "页眉", "有页眉", "全部清空")

    def _clear_footers(self):
        for section in self.doc.sections:
            self._clear_section_footer(section)
        self._record("强制清除", "页脚", "有页脚", "全部清空")

    def _clear_page_numbers(self):
        for section in self.doc.sections:
            self._remove_page_numbers(section)
        self._record("强制清除", "页码", "有页码", "全部移除")

    def _clear_toc(self):
        count = 0
        for para in self.doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""
            if "toc" in style_name:
                count += 1
        if count:
            self._record("强制清除", "目录", f"{count}处", "需手动删除")

    def _clear_hyperlink_colors(self):
        count = 0
        for run in self._iter_all_runs():
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                color = rPr.find(qn('w:color'))
                if color is not None:
                    val = color.get(qn('w:val'))
                    if val and val.upper() not in ('000000', 'AUTO'):
                        color.set(qn('w:val'), '000000')
                        count += 1
                underline = rPr.find(qn('w:u'))
                if underline is not None:
                    underline.set(qn('w:val'), 'none')
        if count:
            self._record("强制清除", "超链接颜色", f"{count}处", "改为黑色无下划线")

    def _clear_char_spacing(self):
        count = 0
        for run in self._iter_all_runs():
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                spacing = rPr.find(qn('w:spacing'))
                if spacing is not None and spacing.get(qn('w:val')):
                    spacing.set(qn('w:val'), '0')
                    count += 1
        if count:
            self._record("强制清除", "字符间距", f"{count}处", "恢复标准")

    def _clear_char_scale(self):
        count = 0
        for run in self._iter_all_runs():
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                w = rPr.find(qn('w:w'))
                if w is not None:
                    val = w.get(qn('w:val'))
                    if val and val != '100':
                        w.set(qn('w:val'), '100')
                        count += 1
        if count:
            self._record("强制清除", "字符缩放", f"{count}处", "恢复100%")

    # ── Page Zones (Testing Reports) ────────────────────────

    def _apply_page_zones_if_needed(self):
        page_zones = self.template.get("page_zones")
        if not page_zones:
            return

        zone_boundaries = {}
        current_zone_id = page_zones[0].get("zone_id") if page_zones else None

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            for zone in page_zones:
                keywords = zone.get("detection_keywords", [])
                for kw in keywords:
                    if kw in text:
                        current_zone_id = zone.get("zone_id")
                        zone_boundaries[i] = current_zone_id
                        break
                if zone_boundaries.get(i) == current_zone_id and i in zone_boundaries:
                    break
            zone_boundaries[i] = current_zone_id

        zone_map = {z.get("zone_id"): z for z in page_zones}
        zone_counts = {}

        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue

            zone_id = zone_boundaries.get(i, current_zone_id)
            zone = zone_map.get(zone_id)
            if not zone:
                continue

            elements = zone.get("elements", [])
            matched = False
            for elem in elements:
                if self._match_zone_element(para, elem, i):
                    self._apply_zone_element(para, elem, zone_id)
                    matched = True
                    break

            if not matched:
                default_elem = self._get_zone_default_element(zone)
                if default_elem and para.text.strip():
                    self._apply_zone_element(para, default_elem, zone_id)

            zone_counts[zone_id] = zone_counts.get(zone_id, 0) + 1

        for zone_id, count in zone_counts.items():
            zone = zone_map.get(zone_id, {})
            zone_name = zone.get("zone_name", zone_id)
            self._record("页面分区", zone_name, "—",
                         f"{count}段已按{zone_name}格式处理")

    def _get_zone_default_element(self, zone):
        """Get the default formatting element for a zone (middle/body element)."""
        elements = zone.get("elements", [])
        for elem in elements:
            pos = elem.get("position", "")
            if pos in ("middle", "line_2_plus", "body"):
                return elem
        default_font = zone.get("default_font", "仿宋")
        default_size = zone.get("default_font_size_pt", 12)
        return {
            "font": default_font,
            "font_size_pt": default_size,
            "bold": False,
            "alignment": "left"
        }

    def _detect_page_zone(self, text, para_idx, total_paras, page_zones, current_zone):
        for zone in page_zones:
            keywords = zone.get("detection_keywords", [])
            position = zone.get("detection_position", "")

            if position == "first_page" and para_idx < 10:
                for kw in keywords:
                    if kw in text:
                        return zone

            if not position:
                for kw in keywords:
                    if kw in text:
                        return zone

        return None

    def _match_zone_element(self, para, elem, idx):
        content_hint = elem.get("content_hint", "")
        if content_hint and content_hint in para.text:
            return True

        position = elem.get("position", "")
        if position == "line_1" and idx == 0:
            return True
        if position == "lines_1_2" and idx <= 1:
            return True

        return False

    def _apply_zone_element(self, para, elem, zone_id=""):
        font = elem.get("font", "仿宋")
        font_size = elem.get("font_size_pt", 12)
        bold = elem.get("bold", False)
        alignment = elem.get("alignment")

        if alignment:
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            if alignment in align_map:
                para.alignment = align_map[alignment]

        for run in para.runs:
            self._format_run(run, font, "Times New Roman", font_size, "000000")
            run.font.bold = bold

    # ── Report Generation ────────────────────────────────────

    def _record(self, location, detail, before, after):
        self.modifications.append({
            "location": location,
            "detail": detail,
            "before": before,
            "after": after,
        })

    def generate_report(self, output_report_path):
        audit_items = [m for m in self.modifications if m["location"] == "文档审查"]
        format_items = [m for m in self.modifications if m["location"] != "文档审查"]

        report = {
            "industry": self.template.get("industry_name", "Unknown"),
            "standard": self.template.get("standard", ""),
            "audit_summary": {
                "issues_found": len(audit_items),
                "details": audit_items,
            },
            "formatting_summary": {
                "total_modifications": len(format_items),
                "modifications": format_items,
            },
            "total_modifications": len(self.modifications),
            "modifications": self.modifications,
        }
        with open(output_report_path, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        return report


def load_template(template_path):
    with open(template_path, "r", encoding="utf-8") as f:
        return json.load(f)


def _sanitize_for_filename(s):
    """Replace characters that are invalid in Windows/macOS filenames."""
    s = re.sub(r'[\\/:*?"<>|]+', '-', s.strip())
    return s.strip('-_. ')


def build_default_output_name(input_path, industry_name):
    """<原文件名>_<行业>_<YYYYMMDD>.docx，保留原文档身份并标明行业与日期。"""
    stem = os.path.splitext(os.path.basename(input_path))[0]
    stem = _sanitize_for_filename(stem)
    industry = _sanitize_for_filename(industry_name or "规范化")
    from datetime import date
    d = date.today().strftime("%Y%m%d")
    return f"{stem}_{industry}_{d}.docx"


def main():
    parser = argparse.ArgumentParser(
        description="DOCX Multi-Industry Document Format Normalizer v2"
    )
    parser.add_argument("--config", help="Path to JSON config file (full config)")
    parser.add_argument("--input", help="Input docx file path")
    parser.add_argument("--template", help="Template JSON file path")
    parser.add_argument("--output", help="Output docx file path")
    parser.add_argument("--report", help="Output report JSON file path")
    args = parser.parse_args()

    if args.config:
        with open(args.config, "r", encoding="utf-8") as f:
            config = json.load(f)
        input_path = config["input_file"]
        template = load_template(config["template_path"]) if "template_path" in config \
                   else config.get("template", {})
        output_path = config.get("output_file") or os.path.join(
            os.path.dirname(os.path.abspath(input_path)),
            build_default_output_name(input_path, template.get("industry_name")))
        report_path = config.get("report_file")
    else:
        if not args.input or not args.template:
            parser.error("--input and --template are required (or use --config)")
        input_path = args.input
        template = load_template(args.template)
        output_path = args.output or os.path.join(
            os.path.dirname(os.path.abspath(input_path)),
            build_default_output_name(input_path, template.get("industry_name")))
        report_path = args.report

    formatter = DocxFormatter({"template": template})
    try:
        formatter.load_document(input_path)
    except RuntimeError as e:
        print(f"ERROR: {e}")
        sys.exit(1)

    print(f"[INFO] Industry: {formatter.industry_name}")
    print(f"[INFO] Input: {input_path}")
    print(f"[INFO] Output: {output_path}")

    modifications = formatter.run(input_path, output_path)
    if report_path:
        formatter.generate_report(report_path)

    print(f"\n[DONE] Processing complete.")
    print(f"  - Total modifications: {len(modifications)}")
    print(f"  - Output file: {output_path}")
    if report_path:
        print(f"  - Report file: {report_path}")
    print(f"\n--- Modification Summary ---")
    for m in modifications:
        print(f"  [{m['location']}] {m['detail']}: {m['before']} -> {m['after']}")


if __name__ == "__main__":
    main()
